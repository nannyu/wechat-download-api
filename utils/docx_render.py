#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (C) 2026 tmwgsicp
# Licensed under the GNU Affero General Public License v3.0
# See LICENSE file in the project root for full license text.
# SPDX-License-Identifier: AGPL-3.0-only
"""
真 .docx 导出。

真 .docx 本质是 ZIP：图片存进 word/media/、**python-docx 按图哈希自动去重**（每篇都有的
分割线/图标只存一份）、再套 ZIP 压缩 → 体积小、质量不降、可编辑。

图片走服务端取图器（export_render.fetch_images_bytes，复用全局并发闸），预抓成 {url:bytes}
映射，再走一遍正文 HTML 用 python-docx 拼段落/图片。

注：本模块的 article 为 dict（单机版 SQLite 行），字段用 .get 访问。
"""
import io
import logging
from datetime import datetime, timezone, timedelta

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

from utils.export_render import fetch_images_bytes, _real_url, _is_wx_image

logger = logging.getLogger(__name__)

# 块级标签：遇到就断段（各自成独立段落）
_BLOCK = {"p", "div", "section", "article", "header", "footer", "figure", "figcaption",
          "blockquote", "li", "ul", "ol", "table", "tr", "td", "h1", "h2", "h3", "h4", "h5", "h6"}
_LINK_COLOR = RGBColor(0x57, 0x6B, 0x95)
_META_COLOR = RGBColor(0x88, 0x88, 0x88)


def _cn_date(pt: int) -> str:
    if not pt:
        return ""
    return datetime.fromtimestamp(pt, timezone(timedelta(hours=8))).strftime("%Y-%m-%d %H:%M")


def _collect_img_urls(articles) -> list:
    urls = []
    for a in articles:
        soup = BeautifulSoup(a.get("content") or "", "html.parser")
        for img in soup.find_all("img"):
            real = _real_url(img.get("src") or img.get("data-src") or "")
            if _is_wx_image(real):
                urls.append(real)
    return urls


def _add_picture(doc: Document, b: bytes):
    """按原始像素宽换算，封顶正文宽 5.8in，不放大小图。"""
    try:
        from PIL import Image
        pw, _ = Image.open(io.BytesIO(b)).size
        w_in = min(5.8, max(1.0, pw / 96.0))
    except Exception:
        w_in = 5.0
    doc.add_picture(io.BytesIO(b), width=Inches(w_in))


def _add_html(doc: Document, html: str, img_map: dict):
    """递归走正文 HTML → docx 段落/图片。维护"当前段落"，块级标签断段、img 单独成段。"""
    soup = BeautifulSoup(html or "", "html.parser")
    state = {"para": None}

    def add_text(text, bold, italic, link):
        text = text.replace("​", "")
        if not text:
            return
        if state["para"] is None:
            state["para"] = doc.add_paragraph()
        run = state["para"].add_run(text)
        run.bold = bold
        run.italic = italic
        if link:
            run.font.color.rgb = _LINK_COLOR
            run.font.underline = True

    def walk(node, bold=False, italic=False, link=False):
        for child in node.children:
            name = getattr(child, "name", None)
            if name is None:  # 文本节点
                add_text(str(child), bold, italic, link)
                continue
            if name in ("script", "style"):
                continue
            if name == "img":
                real = _real_url(child.get("src") or child.get("data-src") or "")
                b = img_map.get(real)
                state["para"] = None  # 断段
                if b:
                    try:
                        _add_picture(doc, b)
                    except Exception as e:
                        logger.debug("docx add_picture 跳过: %s", e)
                state["para"] = None
                continue
            if name == "br":
                if state["para"] is not None:
                    state["para"].add_run().add_break()
                continue
            b2 = bold or name in ("strong", "b")
            i2 = italic or name in ("em", "i")
            l2 = link or name == "a"
            if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                state["para"] = None
                lvl = min(4, int(name[1]))
                try:
                    state["para"] = doc.add_paragraph(style=f"Heading {lvl}")
                except Exception:
                    state["para"] = doc.add_paragraph()
                walk(child, True, i2, l2)
                state["para"] = None
                continue
            if name in _BLOCK:
                state["para"] = None
                walk(child, b2, i2, l2)
                state["para"] = None
                continue
            walk(child, b2, i2, l2)  # 内联标签(span/a/font/sub…)：同段继续

    walk(soup)


def _set_cjk_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    try:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 保证中文渲染
    except Exception:
        pass


async def build_docx(nickname: str, articles: list, note: str,
                     max_side: int = 800, jpeg_q: int = 80) -> bytes:
    """整号文章 → 真 .docx bytes。图片预抓+去重内嵌。articles 为 dict 列表。"""
    img_map = await fetch_images_bytes(_collect_img_urls(articles), max_side, jpeg_q)

    doc = Document()
    _set_cjk_font(doc)

    doc.add_heading(nickname, level=0)
    sub = doc.add_paragraph()
    r = sub.add_run(f"共 {len(articles)} 篇 · 由 WeChat Download API 导出")
    r.font.size = Pt(9)
    r.font.color.rgb = _META_COLOR
    if note:
        rn = doc.add_paragraph().add_run(note)
        rn.font.size = Pt(9)
        rn.font.color.rgb = _META_COLOR

    for a in articles:
        doc.add_heading(a.get("title") or "(无标题)", level=1)
        meta = doc.add_paragraph()
        rm = meta.add_run(f"{a.get('author') or nickname} · {_cn_date(a.get('publish_time') or 0)}")
        rm.font.size = Pt(9)
        rm.font.color.rgb = _META_COLOR
        _add_html(doc, a.get("content") or "", img_map)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
